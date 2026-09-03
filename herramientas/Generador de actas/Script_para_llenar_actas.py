import re
import time
import csv
import unicodedata
from datetime import datetime
import pandas as pd

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from docx import Document

# ========= Configuración =========
USERNAME = "dinatechst@mined.sf"
PASSWORD = "Dinatech2026I!"
URL_BASE = "https://d1i0000001z2ruaq.my.site.com/mined/"

INPUT_IDS_PATH = "Predios_Para_Crear_Actas.xlsx"
EXCEL_OUTPUT_PATH = "Estado_De_Predios.xlsx"
WORD_TEMPLATE = "Acta de relevamiento.docx"
OUTPUT_WORD_DIR = "Actas_Predios"

HEADLESS = False
TIMEOUT = 1            # espera mínima absoluta

# ========= Que equipos entran en el acta =========
# Antes se exigia estado == "Conforme" y todo lo demas se descartaba. Eso dejaba el
# cuadro "Detalle de Red Local" VACIO en predios que si tienen equipos instalados:
# los que estan en "Licencia PNCE", por ejemplo, desaparecian del acta.
#
# La regla es al reves: al acta va todo lo que sigue en el piso, y solo se excluye lo
# que ya no esta fisicamente. Se lista lo que se excluye —no lo que se incluye— para
# que un estado nuevo en Salesforce entre por defecto en vez de desaparecer sin aviso;
# un equipo de mas en el acta se ve y se corrige, uno de menos no se nota.
ESTADOS_EXCLUIDOS = {"robado", "extraviado", "rechazado"}

# Solo para avisar por consola cuando aparece uno que no habiamos visto.
ESTADOS_CONOCIDOS = ESTADOS_EXCLUIDOS | {"conforme", "licencia pnce"}


def _norm_estado(estado):
    """minusculas, sin acentos y sin espacios de mas, para comparar sin sorpresas."""
    txt = unicodedata.normalize("NFD", str(estado or ""))
    txt = "".join(c for c in txt if unicodedata.category(c) != "Mn")
    return " ".join(txt.lower().split())


def equipo_va_al_acta(estado):
    """True si el equipo debe figurar en el cuadro de Detalle de Red Local."""
    return _norm_estado(estado) not in ESTADOS_EXCLUIDOS
WAIT_LABEL = 0.05      # espera casi instantánea para labels comunes
# =================================

PREDIO6_RE = re.compile(r"^\d{6}$")

def _normalize(s) -> str:
    s = str(s).strip()
    s = s.replace("\n", " ").replace("\r", " ")
    s = " ".join(s.split())  # elimina espacios extra
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    return s.lower().replace("ñ", "n")

def crear_driver(headless: bool = False) -> webdriver.Chrome:
    import os
    options = webdriver.ChromeOptions()
    options.page_load_strategy = "eager"
    if headless:
        options.add_argument("--headless=new")
        options.add_argument("--window-size=1920,1080")
    else:
        options.add_argument("--start-maximized")
    prefs = {"profile.managed_default_content_settings.images": 2}
    options.add_experimental_option("prefs", prefs)
    options.add_argument("--disable-notifications")
    options.add_argument("--disable-extensions")
    options.add_argument("--disable-background-networking")
    options.add_argument("--no-sandbox")
    
    # Usar ChromeDriver local si existe
    chromedriver_path = os.path.join(os.path.dirname(__file__), "chromedriver.exe")
    if os.path.exists(chromedriver_path):
        from selenium.webdriver.chrome.service import Service
        service = Service(chromedriver_path)
        driver = webdriver.Chrome(service=service, options=options)
    else:
        # Fallback al ChromeDriver del sistema
        driver = webdriver.Chrome(options=options)
    
    driver.implicitly_wait(0)
    return driver

def login(driver: webdriver.Chrome, username=None, password=None, max_intentos=3):
    for intento in range(1, max_intentos + 1):
        print(f"➡️ Abriendo login... (intento {intento}/{max_intentos})")
        try:
            driver.get(URL_BASE)
            time.sleep(5)  # Esperar más para que cargue Salesforce completamente
            try:
                for f in driver.find_elements(By.TAG_NAME, "iframe"):
                    driver.switch_to.frame(f)
                    if driver.find_elements(By.ID, "username") or driver.find_elements(By.NAME, "username"):
                        break
                    driver.switch_to.default_content()
            except Exception:
                driver.switch_to.default_content()

            def find_first(selectors):
                for by, sel in selectors:
                    els = driver.find_elements(by, sel)
                    if els:
                        return els[0]
                return None

            user = find_first([
                (By.ID, "username"), (By.NAME, "username"),
                (By.CSS_SELECTOR, "input[type='email']"),
                (By.CSS_SELECTOR, "input[type='text']"),
            ])
            pwd = find_first([
                (By.ID, "password"), (By.NAME, "password"),
                (By.CSS_SELECTOR, "input[type='password']"),
            ])

            if not user or not pwd:
                driver.switch_to.default_content()
                print("🔓 No se detectó formulario de login; continúo (probablemente ya autenticado).")
                return True

            # En algunos casos los elementos existen pero no son interactuables
            try:
                try:
                    user.clear()
                except Exception:
                    pass
                try:
                    user.click()
                except Exception:
                    pass
                user.send_keys(username if username is not None else USERNAME)
            except Exception:
                try:
                    driver.execute_script("arguments[0].value = arguments[1]; arguments[0].dispatchEvent(new Event('input'));", user, username if username is not None else USERNAME)
                except Exception:
                    pass

            try:
                try:
                    pwd.clear()
                except Exception:
                    pass
                try:
                    pwd.click()
                except Exception:
                    pass
                pwd.send_keys(password if password is not None else PASSWORD)
            except Exception:
                try:
                    driver.execute_script("arguments[0].value = arguments[1]; arguments[0].dispatchEvent(new Event('input'));", pwd, password if password is not None else PASSWORD)
                except Exception:
                    pass

            btn = find_first([
                (By.ID, "Login"), (By.NAME, "Login"),
                (By.CSS_SELECTOR, "input[type='submit']"),
                (By.CSS_SELECTOR, "button[type='submit']"),
            ])
            if btn: btn.click()
            else:   pwd.submit()

            driver.switch_to.default_content()
            time.sleep(5)  # Esperar a que procese el login
            
            # Verificar si el login fue exitoso
            current_url = driver.current_url
            if "login" not in current_url.lower():
                print("✅ Login exitoso")
                return True
            else:
                print(f"⚠️ Login puede haber fallado, reintentando...")
                time.sleep(3)
                
        except Exception as e:
            print(f"⚠️ Error en intento {intento}: {e}")
            if intento < max_intentos:
                print(f"   Reintentando en 3 segundos...")
                time.sleep(3)
            else:
                print(f"❌ Login falló después de {max_intentos} intentos")
                return False
    
    return False

def leer_csv_robusto(csv_path: str) -> pd.DataFrame:
    """
    Lee una lista de IDs desde un archivo. Si la ruta es Excel (.xls/.xlsx) usa pandas.read_excel,
    si no, intenta autodetectar separador de CSV como antes.
    """
    if str(csv_path).lower().endswith(('.xls', '.xlsx')):
        # Leer hoja por defecto
        df = pd.read_excel(csv_path, dtype=str, keep_default_na=False)
        return df.fillna("")

    # Fallback: CSV
    with open(csv_path, "r", encoding="utf-8-sig", newline="") as f:
        sample = f.read(4096); f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=";,|\t,")
            delimiter = dialect.delimiter
        except Exception:
            delimiter = ","
        reader = csv.reader(f, delimiter=delimiter)
        rows = list(reader)

    if not rows:
        return pd.DataFrame()

    # Simple heuristic: si primera fila parece contener predios
    headerless = all(not isinstance(c, str) for c in rows[0])
    if headerless:
        df = pd.read_csv(csv_path, dtype=str, keep_default_na=False,
                         encoding="utf-8-sig", header=None, delimiter=delimiter)
    else:
        df = pd.read_csv(csv_path, dtype=str, keep_default_na=False,
                         encoding="utf-8-sig", delimiter=delimiter)
    return df.fillna("")

def _texto(elem):
    if not elem:
        return ""
    t = (elem.text or "").strip()
    if t:
        return t
    try:
        v = elem.get_attribute("value")
        return (v or "").strip()
    except Exception:
        return ""

def obtener_por_label(driver, label: str, wait_secs: int = WAIT_LABEL) -> str:
    xpaths = [
        f"//td[normalize-space()='{label}']/following-sibling::td[1]//*[self::div or self::span or self::a or self::p or self::lightning-formatted-text][1]",
        f"//td[normalize-space()='{label}']/following-sibling::td[1]",
        f"//*[normalize-space()='{label}']/following::td[1]",
    ]
    for xp in xpaths:
        try:
            el = WebDriverWait(driver, wait_secs).until(EC.presence_of_element_located((By.XPATH, xp)))
            val = _texto(el)
            if val:
                return val
        except Exception:
            continue
    return ""

def ir_a_registro(driver, record_id: str):
    url = URL_BASE.rstrip("/") + "/" + record_id
    print(f"➡️ Abriendo registro: {record_id}  URL: {url}")
    driver.get(url)
    try:
        WebDriverWait(driver, TIMEOUT).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    except Exception:
        pass

def extraer_campos_predio(driver, record_id: str = None) -> dict:
    out = {
        "Numero de Predio": "",
        "CUE": "",
        "Establecimiento": "",
        "Direccion": "",
        "Localidad": "",
        "Provincia": "",
        "AP_Serial": "",
        "Switch_Serial": "",
        "UTM_Serial": "",
        "Z3_Serial": "",
        "Conex_LAB": "",
        "Estado_Completado": "",  # Nuevo campo para el estado del div
    }
    # --- Extraer lista de APs y Z3 desde el div de equipos ---
    # Ver ESTADOS_EXCLUIDOS arriba: al acta va todo salvo lo que ya no esta en el piso.
    aps = []
    switches = []
    utm_serial = ""
    z3_serial = ""
    equipos_vistos = 0
    try:
        if record_id:
            equipos_div_id = f"{record_id}_00N1I0000061Onc_body"
            equipos_xpath = f"//div[@id='{equipos_div_id}']//tr[contains(@class,'dataRow')]"
            equipos_rows = driver.find_elements(By.XPATH, equipos_xpath)
            equipos_vistos = len(equipos_rows)
            for row in equipos_rows:
                cells = row.find_elements(By.XPATH, "./th | ./td")
                tipo = _texto(cells[4]) if len(cells) > 4 else ""
                sku = _texto(cells[5]) if len(cells) > 5 else ""
                serial = _texto(cells[2]) if len(cells) > 2 else ""
                estado = _texto(cells[3]) if len(cells) > 3 else ""
                if not equipo_va_al_acta(estado):
                    print(f"[DEBUG] Equipo excluido ({estado}): {tipo} serial={serial}")
                    continue
                if _norm_estado(estado) not in ESTADOS_CONOCIDOS:
                    # Se incluye igual, pero queda anotado: si aparece un estado nuevo
                    # que deba excluirse, esto es lo que lo hace visible.
                    print(f"[AVISO] Estado de equipo no visto antes: '{estado}' "
                          f"({tipo} {serial}). Se incluye en el acta.")
                if tipo == "AP":
                    aps.append(serial)
                elif tipo == "Switch":
                    switches.append(serial)
                elif tipo == "UTM":
                    utm_serial = serial
                elif tipo == "Gateway" and "Z3" in sku:
                    z3_serial = serial
    except Exception as e:
        print(f"⚠️ Falla extracción equipos: {e}")
    out["AP_Seriales"] = aps
    out["Switch_Seriales"] = switches
    out["UTM_Serial"] = utm_serial
    out["Z3_Serial"] = z3_serial
    out["Equipos_En_Salesforce"] = equipos_vistos
    out["Equipos_En_Acta"] = len(aps) + len(switches) + (1 if utm_serial else 0) + (1 if z3_serial else 0)

    # Un predio con equipos en Salesforce y el cuadro del acta vacio es la señal de que
    # el filtro se comio algo. Asi salio a campo el acta del 601618, con seis equipos en
    # "Licencia PNCE" y el Detalle de Red Local en blanco, sin que nada lo avisara.
    if equipos_vistos and not out["Equipos_En_Acta"]:
        print(f"⚠️ ATENCION: el predio tiene {equipos_vistos} equipo(s) en Salesforce y "
              f"NINGUNO entró al acta. Revisar los estados antes de usarla.")
    """Extrae los campos solicitados desde la página de un predio abierto en SF.
    Si se proporciona `record_id`, intenta localizar la fila `<tr class='dataRow'>` que contiene ese id
    y toma las celdas por índice (establecimiento=2, direccion=3, localidad=4, provincia=5) tal como en el HTML de ejemplo.
    """

    # --- Extraer desde TODAS las filas <tr class="dataRow"> dentro del div bRelatedList con el record_id ---
    try:
        if record_id:
            # Construir el id completo del div bRelatedList
            div_id = f"{record_id}_00N1I0000061OmT"
            div_xpath = f"//div[contains(@class,'bRelatedList') and @id='{div_id}']"
            div_el = WebDriverWait(driver, 2).until(EC.presence_of_element_located((By.XPATH, div_xpath)))
            # Buscar TODAS las filas dataRow (múltiples escuelas)
            row_els = div_el.find_elements(By.XPATH, ".//tr[contains(@class,'dataRow')]")
            
            # Listas para almacenar datos de múltiples escuelas
            cues_lista = []
            establecimientos_lista = []
            direcciones_lista = []
            localidades_lista = []
            provincias_lista = []
            
            print(f"[DEBUG] Encontradas {len(row_els)} filas dataRow (escuelas) en el predio")
            
            for idx, row_el in enumerate(row_els):
                cells = row_el.find_elements(By.XPATH, "./th | ./td")
                def cell_text(idx):
                    try:
                        return _texto(cells[idx])
                    except Exception:
                        return ""
                
                print(f"[DEBUG] Escuela {idx+1} - Celdas del tr.dataRow:")
                for i, c in enumerate(cells):
                    print(f"  cell[{i}]: '{_texto(c)}'")
                
                # Mapeo según estructura real:
                # cell[1]: CUE
                # cell[2]: Nombre Escuela
                # cell[3]: Calle
                # cell[4]: Ciudad
                # cell[5]: Provincia
                if len(cells) >= 6:
                    cue = cell_text(1)
                    establecimiento = cell_text(2)
                    direccion = cell_text(3)
                    localidad = cell_text(4)
                    provincia = cell_text(5)
                    
                    # Agregar a las listas si no están vacíos
                    if cue:
                        cues_lista.append(cue)
                    if establecimiento:
                        establecimientos_lista.append(establecimiento)
                    if direccion:
                        direcciones_lista.append(direccion)
                    if localidad:
                        localidades_lista.append(localidad)
                    if provincia:
                        provincias_lista.append(provincia)
                    
                    # Numero de Predio: buscar 6-8 dígitos en cell[1] (th) - solo de la primera escuela
                    if idx == 0:
                        candidate_id = cell_text(1)
                        if candidate_id:
                            m = re.search(r"(\d{6,8})", candidate_id)
                            if m:
                                out["Numero de Predio"] = out["Numero de Predio"] or m.group(1)
            
            # Concatenar todos los CUEs con " - "
            if cues_lista:
                out["CUE"] = " - ".join(cues_lista)
                print(f"[DEBUG] CUEs concatenados: {out['CUE']}")
            
            # Para los demás campos, tomar el primer valor o concatenar si es necesario
            if establecimientos_lista:
                out["Establecimiento"] = establecimientos_lista[0]  # Solo el primero
            if direcciones_lista:
                out["Direccion"] = direcciones_lista[0]  # Solo el primero  
            if localidades_lista:
                out["Localidad"] = localidades_lista[0]  # Solo el primero
            if provincias_lista:
                out["Provincia"] = provincias_lista[0]  # Solo el primero
                
    except Exception as e:
        print(f"⚠️ Falla extracción desde div bRelatedList/dataRow: {e}")

    # --- Fallback: extraer desde TODAS las filas <tr class="dataRow"> si no obtuvimos los valores antes ---
    try:
        # Si ya tenemos CUE, establecimiento, etc., no forzamos la extracción
        need_extract = not (out["CUE"] and out["Establecimiento"] and out["Direccion"] and out["Localidad"] and out["Provincia"])
        if need_extract:
            if record_id:
                # Buscar todas las filas que contienen un href con el record_id o el texto del record_id en la columna
                xpath = f"//tr[contains(@class,'dataRow') and (.//a[contains(@href,'/{record_id}')] or .//th//a[normalize-space()='{record_id}'])]"
            else:
                xpath = "//tr[contains(@class,'dataRow')]"
            
            try:
                row_els = driver.find_elements(By.XPATH, xpath)
                if not row_els:
                    # No encontramos filas específicas; intentar cualquier fila dataRow
                    row_els = driver.find_elements(By.XPATH, "//tr[contains(@class,'dataRow')]")
            except Exception:
                row_els = []

            if row_els:
                # Listas para almacenar datos de múltiples escuelas
                cues_lista = []
                establecimientos_lista = []
                direcciones_lista = []
                localidades_lista = []
                provincias_lista = []
                
                print(f"[DEBUG] Fallback: Encontradas {len(row_els)} filas dataRow")
                
                for idx, row_el in enumerate(row_els):
                    # tomar celdas en orden (incluye <th> y <td>)
                    cells = row_el.find_elements(By.XPATH, "./th | ./td")
                    def cell_text(idx):
                        try:
                            return _texto(cells[idx])
                        except Exception:
                            return ""

                    # Log de depuración: mostrar el contenido de todas las celdas
                    print(f"[DEBUG] Fallback Escuela {idx+1} - Celdas del tr.dataRow:")
                    for i, c in enumerate(cells):
                        print(f"  cell[{i}]: '{_texto(c)}'")

                    # Extraer según estructura real:
                    if len(cells) >= 6:
                        cue = cell_text(1)
                        establecimiento = cell_text(2)
                        direccion = cell_text(3)
                        localidad = cell_text(4)
                        provincia = cell_text(5)
                        
                        # Agregar a las listas si no están vacíos
                        if cue:
                            cues_lista.append(cue)
                        if establecimiento:
                            establecimientos_lista.append(establecimiento)
                        if direccion:
                            direcciones_lista.append(direccion)
                        if localidad:
                            localidades_lista.append(localidad)
                        if provincia:
                            provincias_lista.append(provincia)
                        
                        # Numero de Predio: buscar 6 dígitos en cell[1] (th) - solo de la primera escuela
                        if idx == 0:
                            candidate_id = cell_text(1)
                            if candidate_id:
                                m = re.search(r"(\d{6})", candidate_id)
                                if m:
                                    out["Numero de Predio"] = out["Numero de Predio"] or m.group(1)
                
                # Concatenar todos los CUEs con " - " solo si no teníamos CUE antes
                if cues_lista and not out["CUE"]:
                    out["CUE"] = " - ".join(cues_lista)
                    print(f"[DEBUG] Fallback - CUEs concatenados: {out['CUE']}")
                
                # Para los demás campos, tomar el primer valor solo si no los teníamos antes
                if establecimientos_lista and not out["Establecimiento"]:
                    out["Establecimiento"] = establecimientos_lista[0]
                if direcciones_lista and not out["Direccion"]:
                    out["Direccion"] = direcciones_lista[0]
                if localidades_lista and not out["Localidad"]:
                    out["Localidad"] = localidades_lista[0]
                if provincias_lista and not out["Provincia"]:
                    out["Provincia"] = provincias_lista[0]
                    
    except Exception as e:
        # No fatal — mantenemos lo que ya tenemos
        print(f"⚠️ Falla extracción fallback desde filas dataRow: {e}")

    # --- Extraer Numero de Predio desde <h2 class="pageDescription"> ---
    try:
        h2_el = WebDriverWait(driver, 1).until(EC.presence_of_element_located((By.XPATH, "//h2[contains(@class,'pageDescription')]")))
        h2_text = (h2_el.text or "").strip()
        m = re.search(r"(\d{6})", h2_text)
        if m:
            out["Numero de Predio"] = m.group(1)
            print(f"[DEBUG] Numero de Predio desde pageDescription: {out['Numero de Predio']}")
    except Exception as e:
        print(f"⚠️ No se pudo extraer Numero de Predio desde pageDescription: {e}")

    # --- Extraer Estado Completado desde div específico ---
    try:
        div_completado = driver.find_element(By.ID, "00N1I0000061OmB_ileinner")
        texto_completado = (div_completado.text or "").strip()
        out["Estado_Completado"] = texto_completado
        print(f"[DEBUG] Estado Completado encontrado: '{texto_completado}'")
    except Exception as e:
        print(f"⚠️ No se pudo extraer Estado Completado desde div: {e}")
        out["Estado_Completado"] = ""

    print("   -> Campos extraídos:", out)
    return out

def escribir_celda_preservando_formato(cell, valor: str):
    """Escribe texto en una celda con formato Calibri 9pt y sangría 5pt."""
    from docx.shared import Pt
    
    if not cell.paragraphs:
        cell.text = str(valor)
        return
    
    # Tomar el primer párrafo
    para = cell.paragraphs[0]
    
    # Limpiar todos los runs del párrafo
    for run in para.runs:
        run.clear()
    
    # Si no hay runs, crear uno nuevo
    if not para.runs:
        run = para.add_run(str(valor))
    else:
        # Usar el primer run existente
        para.runs[0].text = str(valor)
        run = para.runs[0]
    
    # Aplicar formato Calibri 9pt (igual que la plantilla)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    
    # Aplicar sangría 5pt
    para.paragraph_format.left_indent = Pt(5)
    
    # Limpiar párrafos adicionales si existen
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)

def rellenar_word(template_path: str, out_data: dict, out_path: str):
    doc = Document(template_path)
    if not doc.tables:
        raise RuntimeError("La plantilla Word no contiene tablas.")
    
    # Log simplificado de tablas relevantes (actualizado para plantilla nueva con 9 tablas)
    print("[LOG] Estructura de tablas relevantes:")
    if len(doc.tables) > 0:
        print(f"[LOG] TABLA 1 (datos generales, índice 0): {len(doc.tables[0].rows)} filas")
    if len(doc.tables) > 4:
        print(f"[LOG] TABLA 5 (APs 1-22, índice 4): {len(doc.tables[4].rows)} filas")
    if len(doc.tables) > 6:
        print(f"[LOG] TABLA 7 (APs 23-44, índice 6): {len(doc.tables[6].rows)} filas")
    if len(doc.tables) > 7:
        print(f"[LOG] TABLA 8 (conectividad, índice 7): {len(doc.tables[7].rows)} filas")
    if len(doc.tables) > 8:
        print(f"[LOG] TABLA 9 (conexión LAB, índice 8): {len(doc.tables[8].rows)} filas")
    
    # Escritura de datos generales en hoja 1 (tabla 0) y hoja 5 (tabla 4)
    def find_row_by_label(label, table):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text and _normalize(label) in _normalize(cell.text):
                    return r_idx, c_idx
        return None, None
    
    # Función auxiliar para normalizar etiquetas de AP (maneja "AP Nº10" sin espacio)
    def normalizar_ap_label(label):
        return _normalize(label).replace(" ", "")
    
    mapping = [
        ("Predio", out_data.get("Numero de Predio", "")),
        ("CUE", out_data.get("CUE", "")),
        ("INCIDENCIA ASOCIADA", out_data.get("Numero_Incidencia", "")),
        ("Establecimiento", out_data.get("Establecimiento", "")),
        ("Dirección", out_data.get("Direccion", "")),
        ("Localidad", out_data.get("Localidad", "")),
        ("Provincia", out_data.get("Provincia", "")),
        ("PROVEEDOR SERVICIO TECNICO", "DINATECH ST"),
    ]
    # Hoja 1 (tabla 0)
    tbl1 = doc.tables[0]
    print("[DEBUG] Antes de escribir hoja 1:")
    for r_idx, row in enumerate(tbl1.rows):
        for c_idx, cell in enumerate(row.cells):
            print(f"  fila {r_idx+1}, col {c_idx+1}: '{cell.text.strip()}'")
    for label, val in mapping:
        r_idx, c_idx = find_row_by_label(label, tbl1)
        if r_idx is not None:
            if c_idx + 1 < len(tbl1.rows[r_idx].cells):
                next_cell = tbl1.rows[r_idx].cells[c_idx + 1]
                if not next_cell.text.strip() or _normalize(next_cell.text) == _normalize(val):
                    print(f"[DEBUG] Escribiendo '{label}' en hoja 1, fila {r_idx+1}, col {c_idx+2}: '{val}'")
                    escribir_celda_preservando_formato(next_cell, val)
                else:
                    print(f"⚠️ Celda destino para '{label}' en hoja 1 no está vacía: '{next_cell.text.strip()}'")
            else:
                print(f"⚠️ No hay celda destino para '{label}' en hoja 1.")
        else:
            print(f"⚠️ No encontré la fila con label '{label}' en hoja 1.")
    print("[DEBUG] Después de escribir hoja 1:")
    for r_idx, row in enumerate(tbl1.rows):
        for c_idx, cell in enumerate(row.cells):
            print(f"  fila {r_idx+1}, col {c_idx+1}: '{cell.text.strip()}'")
    # Tabla 8 (índice 7) - RELEVAMIENTO DE LA CONECTIVIDAD A INTERNET (actualizado para plantilla nueva con 9 tablas)
    if len(doc.tables) > 7:
        tbl_conectividad = doc.tables[7]
        print("[DEBUG] Escribiendo datos generales en tabla 8 (conectividad):")
        # Mapeo específico para tabla 9 sin INCIDENCIA ASOCIADA, incluye PROVEEDOR SERVICIO TECNICO
        mapping_tbl_conectividad = [
            ("Predio", out_data.get("Numero de Predio", "")),
            ("CUE", out_data.get("CUE", "")),
            ("Establecimiento", out_data.get("Establecimiento", "")),
            ("Dirección", out_data.get("Direccion", "")),
            ("Localidad", out_data.get("Localidad", "")),
            ("Provincia", out_data.get("Provincia", "")),
            ("PROVEEDOR", "DINATECH ST"),
        ]
        for label, val in mapping_tbl_conectividad:
            r_idx, c_idx = find_row_by_label(label, tbl_conectividad)
            if r_idx is not None:
                if c_idx + 1 < len(tbl_conectividad.rows[r_idx].cells):
                    next_cell = tbl_conectividad.rows[r_idx].cells[c_idx + 1]
                    if not next_cell.text.strip() or _normalize(next_cell.text) == _normalize(val):
                        print(f"[DEBUG] Escribiendo '{label}' en tabla 8, fila {r_idx+1}, col {c_idx+2}: '{val}'")
                        escribir_celda_preservando_formato(next_cell, val)
                    else:
                        print(f"⚠️ Celda destino para '{label}' en tabla 8 no está vacía: '{next_cell.text.strip()}'")
                else:
                    print(f"⚠️ No hay celda destino para '{label}' en tabla 8.")
            else:
                print(f"⚠️ No encontré la fila con label '{label}' en tabla 8.")
        # Lógica antigua para escribir en tabla 5 eliminada - ahora se usa tabla 8
    
    # Nueva lógica: escribir equipos en tabla 5 (índice 4) con capacidades ampliadas
    if len(doc.tables) > 4:
        ap_tbl = doc.tables[4]  # Tabla 5 (índice 4) - la que tiene los APs del 1 al 22
        
        # APs: escribimos hasta 22 APs en tabla 5 (columna izquierda: 1-14, columna derecha: 15-22)
        aps = out_data.get("AP_Seriales", [])
        
        # Columna izquierda: AP Nº 1-14 (filas 2-15)
        for i in range(min(14, len(aps))):
            label = f"AP Nº {i+1}"  # Con espacio para todos los APs
            for r, row in enumerate(ap_tbl.rows):
                # Comparar sin espacios para manejar "AP Nº10" vs "AP Nº 10"
                if len(row.cells) > 1 and normalizar_ap_label(row.cells[0].text) == normalizar_ap_label(label):
                    ap_val = aps[i] if i < len(aps) else ""
                    escribir_celda_preservando_formato(ap_tbl.rows[r].cells[1], ap_val)
                    print(f"[DEBUG] Escribiendo AP {i+1}: '{ap_val}' en tabla 5, columna izquierda, fila {r+1}")
                    break
        
        # Columna derecha: AP Nº 15-22 (filas 2-9)
        for i in range(14, min(22, len(aps))):
            label = f"AP Nº {i+1}"
            for r, row in enumerate(ap_tbl.rows):
                if len(row.cells) > 3 and normalizar_ap_label(row.cells[2].text) == normalizar_ap_label(label):
                    ap_val = aps[i] if i < len(aps) else ""
                    escribir_celda_preservando_formato(ap_tbl.rows[r].cells[3], ap_val)
                    print(f"[DEBUG] Escribiendo AP {i+1}: '{ap_val}' en tabla 5, columna derecha, fila {r+1}")
                    break
        
        # Switches: ahora escribimos hasta 3 switches en tabla 5
        switches = out_data.get("Switch_Seriales", [])
        for i in range(min(3, len(switches))):  # Cambiado de 4 a 3 switches
            label = f"Switch Nº {i+1}"
            for r, row in enumerate(ap_tbl.rows):
                if len(row.cells) > 3 and _normalize(row.cells[2].text) == _normalize(label):
                    sw_val = switches[i] if i < len(switches) else ""
                    escribir_celda_preservando_formato(ap_tbl.rows[r].cells[3], sw_val)
                    print(f"[DEBUG] Escribiendo Switch {i+1}: '{sw_val}' en tabla 5, fila {r+1}")
                    break
        
        # UTM: escribir en tabla 5
        utm_val = out_data.get("UTM_Serial", "")
        for r, row in enumerate(ap_tbl.rows):
            if len(row.cells) > 3 and _normalize(row.cells[2].text) == _normalize("UTM"):
                escribir_celda_preservando_formato(ap_tbl.rows[r].cells[3], utm_val)
                print(f"[DEBUG] Escribiendo UTM: '{utm_val}' en tabla 5, fila {r+1}")
                break
        
        # Z3: escribir en tabla 5
        z3_val = out_data.get("Z3_Serial", "")
        for r, row in enumerate(ap_tbl.rows):
            if len(row.cells) > 3 and _normalize(row.cells[2].text) == _normalize("Z3"):
                escribir_celda_preservando_formato(ap_tbl.rows[r].cells[3], z3_val)
                print(f"[DEBUG] Escribiendo Z3: '{z3_val}' en tabla 5, fila {r+1}")
                break

    # APs extra: si hay más de 22 APs, continuar en tabla 7 (índice 6) - APs 23-44
    if len(doc.tables) > 6 and len(aps) > 22:
        ap_tbl_extra = doc.tables[6]  # Tabla 7 (índice 6)
        
        # Columna izquierda: AP Nº 23-36 (filas 2-15)
        for i in range(22, min(36, len(aps))):  # APs 23-36
            label = f"AP Nº {i+1}"
            for r, row in enumerate(ap_tbl_extra.rows):
                if len(row.cells) > 1 and normalizar_ap_label(row.cells[0].text) == normalizar_ap_label(label):
                    ap_val = aps[i] if i < len(aps) else ""
                    escribir_celda_preservando_formato(ap_tbl_extra.rows[r].cells[1], ap_val)
                    print(f"[DEBUG] Escribiendo AP {i+1}: '{ap_val}' en tabla 7, columna izquierda, fila {r+1}")
                    break
        
        # Columna derecha: AP Nº 37-44 (filas 2-9)
        for i in range(36, min(44, len(aps))):  # APs 37-44
            label = f"AP Nº {i+1}"
            for r, row in enumerate(ap_tbl_extra.rows):
                if len(row.cells) > 3 and normalizar_ap_label(row.cells[2].text) == normalizar_ap_label(label):
                    ap_val = aps[i] if i < len(aps) else ""
                    escribir_celda_preservando_formato(ap_tbl_extra.rows[r].cells[3], ap_val)
                    print(f"[DEBUG] Escribiendo AP {i+1}: '{ap_val}' en tabla 7, columna derecha, fila {r+1}")
                    break

    # Switches extra: si hay más de 3 switches, continuar en tabla 7 (índice 6) - Switches 4 en adelante
    if len(doc.tables) > 6 and len(switches) > 3:
        switches_tbl_extra = doc.tables[6]  # Tabla 7 (índice 6)
        
        # Escribir switches adicionales en la tabla 7 donde haya espacio disponible
        # Buscar etiquetas "Switch Nº 4", "Switch Nº 5", etc.
        for i in range(3, min(6, len(switches))):  # Switches del 4 al 6
            label = f"Switch Nº {i+1}"
            for r, row in enumerate(switches_tbl_extra.rows):
                # Buscar en ambas columnas (izquierda y derecha)
                if len(row.cells) > 1 and _normalize(row.cells[0].text) == _normalize(label):
                    sw_val = switches[i] if i < len(switches) else ""
                    escribir_celda_preservando_formato(switches_tbl_extra.rows[r].cells[1], sw_val)
                    print(f"[DEBUG] Escribiendo Switch {i+1}: '{sw_val}' en tabla 7, columna izquierda, fila {r+1}")
                    break
                elif len(row.cells) > 3 and _normalize(row.cells[2].text) == _normalize(label):
                    sw_val = switches[i] if i < len(switches) else ""
                    escribir_celda_preservando_formato(switches_tbl_extra.rows[r].cells[3], sw_val)
                    print(f"[DEBUG] Escribiendo Switch {i+1}: '{sw_val}' en tabla 7, columna derecha, fila {r+1}")
                    break

    # Escribir estado completado en tabla 9 (índice 8), fila 2, columna 3
    estado_completado = out_data.get("Estado_Completado", "")
    if len(doc.tables) > 8:  # Verificar que existe la tabla 9 (índice 8)
        tabla_lab = doc.tables[8]
        try:
            # Verificar que existe fila 2 y columna 3 (CONEXIÓN LAB PROYECTO PNCE)
            if len(tabla_lab.rows) > 1 and len(tabla_lab.rows[1].cells) > 2:
                if estado_completado.lower() == "completada":
                    valor_escribir = "SI"
                elif estado_completado.lower() == "no aplica":
                    valor_escribir = "NO"
                else:
                    valor_escribir = ""  # Si no es ninguno de los dos valores esperados
                
                if valor_escribir:
                    escribir_celda_preservando_formato(tabla_lab.rows[1].cells[2], valor_escribir)
                    print(f"[DEBUG] Estado '{estado_completado}' -> Escribiendo '{valor_escribir}' en tabla 9, fila 2, columna 3 (CONEXIÓN LAB)")
                else:
                    print(f"[DEBUG] Estado '{estado_completado}' no reconocido, no se escribe nada en tabla 9")
            else:
                print("⚠️ La tabla 9 no tiene suficientes filas o columnas")
        except Exception as e:
            print(f"⚠️ Error escribiendo en tabla 9: {e}")
    else:
        print("⚠️ No existe la tabla 9 en el documento Word")

    doc.save(out_path)
    print(f"✅ Guardado Word: {out_path}")

def procesar(driver, df: pd.DataFrame) -> list[dict]:
    registros = []
    headerless = all(not isinstance(c, str) for c in df.columns)

    for i, row in df.iterrows():
        # record_id como primera columna
        try:
            record_id = str(row.values[0]).strip()
        except Exception:
            record_id = ""
        
        # Número de incidencia desde columna 3 (índice 2)
        try:
            numero_incidencia = str(row.values[2]).strip() if len(row.values) > 2 else ""
        except Exception:
            numero_incidencia = ""

        if not record_id:
            print("❌ Fila sin record_id — salto.")
            continue

        try:
            ir_a_registro(driver, record_id)
            data = extraer_campos_predio(driver, record_id=record_id)
            # Agregar número de incidencia al diccionario de datos
            data["Numero_Incidencia"] = numero_incidencia
        except Exception as e:
            print(f"⚠️ Error extrayendo {record_id}: {e}")
            data = {"Numero_Incidencia": numero_incidencia}

        # Guardar Word por predio
        import os
        os.makedirs(OUTPUT_WORD_DIR, exist_ok=True)
        out_name = os.path.join(OUTPUT_WORD_DIR, f"Acta_{data.get('Numero de Predio', record_id)}.docx")
        try:
            rellenar_word(WORD_TEMPLATE, data, out_name)
        except Exception as e:
            print(f"⚠️ Error rellenando Word para {record_id}: {e}")

        registros.append({"RecordID": record_id, **data})
    return registros

def guardar_excel(registros: list[dict], ruta: str):
    df = pd.DataFrame(registros)
    try:
        df.to_excel(ruta, index=False)
        print(f"\n✅ Archivo Excel guardado como '{ruta}'")
    except PermissionError:
        ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        alt = ruta.replace(".xlsx", f"_{ts}.xlsx")
        df.to_excel(alt, index=False)
        print(f"\n⚠️ '{ruta}' estaba en uso. Guardé como '{alt}'")

def main():
    print("📄 Leyendo CSV (detección de separador y cabecera)...")
    df = leer_csv_robusto(INPUT_IDS_PATH)
    print(f"   Forma: {df.shape}")

    driver = crear_driver(headless=HEADLESS)
    try:
        login(driver)
        registros = procesar(driver, df)
        guardar_excel(registros, EXCEL_OUTPUT_PATH)
    finally:
        try:
            driver.quit()
        except Exception:
            pass

if __name__ == "__main__":
    main()
